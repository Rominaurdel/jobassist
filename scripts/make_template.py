#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Crée le template DOCX pour docxtpl
Usage: python make_template.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)

# Contenu unique avec variable docxtpl
content_para = doc.add_paragraph()
content_run = content_para.add_run("{{cv_content}}")
content_run.font.size = Pt(9)
content_run.font.color.rgb = RGBColor(51, 51, 51)
content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Sauvegarder
doc.save("CV_Template_Simple.docx")
print("✅ Template créé: CV_Template_Simple.docx")
